from __future__ import annotations

from pathlib import Path
import math
import re

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2] / "docs" / "si"
OUT_PATH = ROOT / "supplementary_information_results123.docx"
LEGACY_OUT_PATH = ROOT / "result3_si_analysis_figures.docx"

TABLE_DIR = ROOT / "tables"
FIG_DIR = ROOT / "figures"
SUMMARY_DIR = ROOT / "summaries"
MAP_DIR = Path(__file__).resolve().parents[2] / "inputs" / "auxiliary" / "maps"

SDG_LABELS = {
    1: "No Poverty",
    2: "Zero Hunger",
    3: "Good Health and Well-being",
    4: "Quality Education",
    5: "Gender Equality",
    6: "Clean Water and Sanitation",
    7: "Affordable and Clean Energy",
    8: "Decent Work and Economic Growth",
    9: "Industry, Innovation and Infrastructure",
    10: "Reduced Inequalities",
    11: "Sustainable Cities and Communities",
    12: "Responsible Consumption and Production",
    13: "Climate Action",
    14: "Life Below Water",
    15: "Life on Land",
    16: "Peace, Justice and Strong Institutions",
    17: "Partnerships for the Goals",
}

FIGURE_NOTES = {
    "Figure_SB1_sdg_goal_phase_stacked.png": (
        "Figure SB1. SDG goal composition over time.",
        "Interpretation: this figure visualizes how the annual SDG portfolio shifts across goals rather than reporting a single aggregate total. "
        "It is useful for checking whether broad output growth coincides with diversification in SDG attention.",
        "Source: existing saved plot generated from goal-level yearly count and share tables retained in overton_sdg_analysis.xlsx.",
    ),
    "Figure_SB2_sdg_goal_inequality_phase3_vs_phase4.png": (
        "Figure SB2. SDG inequality comparison between Phase 3 and Phase 4.",
        "Interpretation: this figure compares concentration metrics across mature and pre-expansion phases, showing whether SDG attention became more or less concentrated under later policy expansion.",
        "Source: existing saved plot generated from phase-specific concentration outputs already stored in the project.",
    ),
    "Figure_SB3_scatter_oadr_vs_index.png": (
        "Figure SB3. Country-level association between OADR and the validated final v2 ageing-energy index.",
        "Interpretation: this plot is the visual companion to the pooled and fixed-effects regressions, and it shows whether more aged countries systematically align more strongly with the ageing-energy SDG profile.",
        "Source: existing saved ageing-energy story figure based on country-level OADR and validated final v2 index outputs.",
    ),
    "Figure_SB4_scatter_index_vs_logimpact.png": (
        "Figure SB4. Country-level association between the validated final v2 ageing-energy index and policy impact.",
        "Interpretation: this plot evaluates whether stronger ageing-energy alignment is associated with higher policy citation impact, while allowing scale differences to remain visible.",
        "Source: existing saved ageing-energy story figure based on country-level impact and final v2 index outputs.",
    ),
    "Figure_SB5_scatter_index_vs_efficiency.png": (
        "Figure SB5. Country-level association between the validated final v2 ageing-energy index and policy efficiency.",
        "Interpretation: this figure assesses whether countries with stronger thematic alignment also translate output into citations more efficiently.",
        "Source: existing saved ageing-energy story figure derived from efficiency and validated final v2 index outputs.",
    ),
    "Figure_SB6_inequality_timeseries.png": (
        "Figure SB6. Time-series of inequality metrics for the validated final v2 ageing-energy index.",
        "Interpretation: this figure tracks whether ageing-energy alignment became more concentrated or more diffused across countries over time.",
        "Source: existing saved inequality figure from ageing_energy_story_analysis.xlsx outputs.",
    ),
    "Figure_SC1_topic_elasticity_ranking.png": (
        "Figure SC1. Policy attention elasticity ranking across the six Result 3 taxonomy dimensions.",
        "Interpretation: this figure ranks the elasticity estimates from Step 8 and makes it easy to compare household-level mechanisms against broader macro themes under rising ageing pressure.",
        "Source: existing saved Step 8 plot built directly from step8_pae_ranked.csv without rerunning analysis.",
    ),
    "Figure_SC2_centrality_comparison.png": (
        "Figure SC2. Cross-phase centrality comparison of key macro and mechanism-related terms.",
        "Interpretation: this figure condenses the phase-wise network evidence into a cross-phase comparison of how central household mechanisms remain relative to macro terms.",
        "Source: existing saved Step 9 comparison plot generated from step9_centrality_comparison.csv.",
    ),
    "Figure_SC3_phase1_network.png": (
        "Figure SC3. Phase 1 keyword co-occurrence network.",
        "Interpretation: this network shows the sparse early discursive core before later policy expansion.",
        "Source: existing saved Step 9 network plot from Phase 1 keyword co-occurrence metrics.",
    ),
    "Figure_SC4_phase2_network.png": (
        "Figure SC4. Phase 2 keyword co-occurrence network.",
        "Interpretation: this network shows the initial institutional broadening of the discourse while remaining structurally limited.",
        "Source: existing saved Step 9 network plot from Phase 2 keyword co-occurrence metrics.",
    ),
    "Figure_SC5_phase3_network.png": (
        "Figure SC5. Phase 3 keyword co-occurrence network.",
        "Interpretation: this network captures the transition period in which macro keywords still dominate the core while mechanism terms appear but remain peripheral.",
        "Source: existing saved Step 9 network plot from Phase 3 keyword co-occurrence metrics.",
    ),
    "Figure_SC6_phase4_network.png": (
        "Figure SC6. Phase 4 keyword co-occurrence network.",
        "Interpretation: this network shows the expanded semantic field under rapid expansion, while still allowing direct inspection of whether key household-level terms move toward or remain outside the discursive core.",
        "Source: existing saved Step 9 network plot from Phase 4 keyword co-occurrence metrics.",
    ),
    "Figure_SC7_KL_divergence_trend.png": (
        "Figure SC7. KL divergence trend over time under the equal-weight six-category reference distribution.",
        "Interpretation: this figure provides a supplementary diagnostic of mechanism mismatch by comparing the observed topic distribution with a balanced six-category reference distribution.",
        "Source: existing saved Step 10 plot generated from step10_kl_timeseries.csv with no re-estimation.",
    ),
}


def set_default_font(document: Document, name: str = "Arial", size: int = 10) -> None:
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = document.styles[style_name]
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        if style_name == "Title":
            style.font.size = Pt(size + 7)
        elif style_name == "Heading 1":
            style.font.size = Pt(size + 4)
        elif style_name == "Heading 2":
            style.font.size = Pt(size + 2)
        elif style_name == "Heading 3":
            style.font.size = Pt(size + 1)
        else:
            style.font.size = Pt(size)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(8)
    run.italic = True


def format_value(x):
    if pd.isna(x):
        return ""
    if isinstance(x, (int,)) or (isinstance(x, float) and x.is_integer()):
        return f"{int(x)}"
    if isinstance(x, float):
        ax = abs(x)
        if ax >= 1000:
            return f"{x:,.1f}"
        if ax >= 1:
            return f"{x:.4f}".rstrip("0").rstrip(".")
        return f"{x:.6f}".rstrip("0").rstrip(".")
    return str(x)


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)


def add_table_from_df(doc: Document, title: str, explanation: str, source_note: str, df: pd.DataFrame) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(explanation)
    add_source_note(doc, source_note)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(section)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], str(col), bold=True, size=8)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], format_value(val), size=7)

    doc.add_paragraph()
    next_section = doc.add_section(WD_SECTION.NEW_PAGE)
    next_section.orientation = WD_ORIENT.PORTRAIT
    next_section.page_width, next_section.page_height = next_section.page_height, next_section.page_width
    next_section.left_margin = Inches(0.8)
    next_section.right_margin = Inches(0.8)
    next_section.top_margin = Inches(0.7)
    next_section.bottom_margin = Inches(0.7)


def insert_figure(doc: Document, fig_path: Path, width: float = 6.2) -> None:
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Inches(width))


def add_figure_block(doc: Document, fig_name: str) -> None:
    fig_path = FIG_DIR / fig_name
    if not fig_path.exists():
        return
    caption, interpretation, source = FIGURE_NOTES[fig_name]
    doc.add_heading(caption.split(".")[0], level=2)
    insert_figure(doc, fig_path, width=6.5)
    add_caption(doc, caption)
    doc.add_paragraph(interpretation)
    add_source_note(doc, source + f" File retained at {fig_path.as_posix()}.")


def parse_sdg_map_name(path: Path) -> tuple[int, str] | None:
    m = re.match(r"sdg(\d+)_phase([34])\.png", path.name)
    if not m:
        return None
    return int(m.group(1)), f"Phase {m.group(2)}"


def add_sdg_map_appendix(doc: Document) -> None:
    phase3 = {parse_sdg_map_name(p)[0]: p for p in sorted(MAP_DIR.glob("sdg*_phase3.png")) if parse_sdg_map_name(p)}
    phase4 = {parse_sdg_map_name(p)[0]: p for p in sorted(MAP_DIR.glob("sdg*_phase4.png")) if parse_sdg_map_name(p)}
    if not phase3 and not phase4:
        return

    doc.add_page_break()
    doc.add_heading("Supplementary SDG Goal Maps", level=1)
    doc.add_paragraph(
        "The following SDG goal maps are retained as supplementary visual references from the existing project outputs. "
        "They are not newly generated here. Each goal is shown with the saved Phase 3 and Phase 4 map where available."
    )

    for sdg in sorted(set(phase3) | set(phase4)):
        label = SDG_LABELS.get(sdg, f"SDG {sdg}")
        doc.add_heading(f"SDG {sdg}. {label}", level=2)
        t = doc.add_table(rows=1, cols=2)
        t.autofit = True
        row = t.rows[0].cells
        for idx, (phase, mapping) in enumerate([("Phase 3", phase3), ("Phase 4", phase4)]):
            cell = row[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(phase)
            run.bold = True
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(9)
            img_path = mapping.get(sdg)
            if img_path and img_path.exists():
                cell.add_paragraph().add_run().add_picture(str(img_path), width=Inches(3.0))
            else:
                cell.add_paragraph("Image not available in current project tree.")
        add_caption(doc, f"Supplementary SDG map pair for SDG {sdg} ({label}).")
        doc.add_paragraph(
            f"Interpretation: these retained maps provide a visual complement to the SDG diffusion and concentration tables by showing how the geographical footprint of SDG {sdg} changes between Phase 3 and Phase 4."
        )
        add_source_note(
            doc,
            f"Source: saved map outputs under {MAP_DIR.as_posix()}, derived previously from the SDG goal-by-country phase aggregation used elsewhere in Result 2.",
        )


def build_document() -> Document:
    doc = Document()
    set_default_font(doc, "Arial", 10)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Information for Results 1–3")
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(17)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Compiled from existing validated project outputs only; no analyses were rerun during this assembly step."
    )
    r.italic = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(10)

    # Note 1
    doc.add_heading("Supplementary Note 1. Temporal segmentation and structural evolution", level=1)
    doc.add_paragraph(
        "This note assembles the saved quantitative materials supporting Result 1. "
        "The retained four-phase structure is the segmentation used throughout the current Figure 1 package. "
        "The saved outputs identify transition years at 1948, 1972, and 1994, corresponding to end-of-phase breakpoints at 1947, 1971, and 1993. "
        "A complete candidate-model BIC comparison table is not present among the saved outputs, so this SI documents the retained breakpoint structure and the validated phase-level quantitative characteristics rather than reconstructing omitted model-selection runs."
    )
    doc.add_paragraph(
        SUMMARY_DIR.joinpath("Result1_figure1_quantitative_interpretation.txt").read_text(encoding="utf-8")
    )

    # Note 2
    doc.add_heading("Supplementary Note 2. SDG alignment, diffusion and concentration", level=1)
    doc.add_paragraph(
        "This note consolidates the non-archived SDG outputs currently retained in the project tree. "
        "It covers goal-level diffusion breadth, diffusion intensity, diffusion speed, concentration metrics, "
        "the validated final v2 ageing–energy target weights, and the saved regression outputs linking demographic ageing, policy scale, impact, efficiency, and thematic alignment."
    )
    doc.add_paragraph(
        SUMMARY_DIR.joinpath("Result2_overton_sdg_methods_note.txt").read_text(encoding="utf-8")
    )

    # Note 3
    doc.add_heading("Supplementary Note 3. Blind spots in policy attention", level=1)
    doc.add_heading("3.1 Taxonomy and thematic classification", level=2)
    doc.add_paragraph(
        "The clean keyword taxonomy retained for Result 3 excludes low-relevance keywords from the classified set and avoids creating a generic governance category. "
        "The active categories are welfare and care, health and vulnerability, housing and thermal comfort, energy transition and efficiency, household behaviour and adoption, and affordability and income constraints."
    )
    doc.add_heading("3.2 Topic intensity and temporal composition", level=2)
    doc.add_paragraph(
        "TopicShare values are computed only over the classified keyword denominator rather than over the full token set. "
        "This makes the temporal composition table interpretable as a within-taxonomy allocation rather than a general corpus share."
    )
    doc.add_heading("3.3 Policy Attention Elasticity", level=2)
    doc.add_paragraph(SUMMARY_DIR.joinpath("Result3_step8_pae_summary.md").read_text(encoding="utf-8"))
    doc.add_heading("3.4 Network path dependency and semantic centrality", level=2)
    doc.add_paragraph(SUMMARY_DIR.joinpath("Result3_step9_network_summary.md").read_text(encoding="utf-8"))
    doc.add_heading("3.5 Mechanism mismatch quantified by KL divergence", level=2)
    doc.add_paragraph(SUMMARY_DIR.joinpath("Result3_step10_kl_summary.md").read_text(encoding="utf-8"))

    # Main SI figures with names/explanations/sources
    doc.add_page_break()
    doc.add_heading("Supplementary Figures for Results 2 and 3", level=1)
    doc.add_paragraph(
        "Each figure below is retained from the existing project outputs. "
        "The accompanying text states what the figure shows, how it should be read, and which saved result file it comes from."
    )
    for fig_name in [
        "Figure_SB1_sdg_goal_phase_stacked.png",
        "Figure_SB2_sdg_goal_inequality_phase3_vs_phase4.png",
        "Figure_SB3_scatter_oadr_vs_index.png",
        "Figure_SB4_scatter_index_vs_logimpact.png",
        "Figure_SB5_scatter_index_vs_efficiency.png",
        "Figure_SB6_inequality_timeseries.png",
        "Figure_SC1_topic_elasticity_ranking.png",
        "Figure_SC2_centrality_comparison.png",
        "Figure_SC3_phase1_network.png",
        "Figure_SC4_phase2_network.png",
        "Figure_SC5_phase3_network.png",
        "Figure_SC6_phase4_network.png",
        "Figure_SC7_KL_divergence_trend.png",
    ]:
        add_figure_block(doc, fig_name)

    add_sdg_map_appendix(doc)

    # Tables as formatted word tables
    table_meta = [
        (
            "Table S1. Retained phase breakpoints and segmentation support",
            "This table records the retained breakpoint structure used in the validated Figure 1 package and notes the absence of a separately saved candidate-model BIC comparison table.",
            f"Derived from existing segmentation outputs summarized in {TABLE_DIR.joinpath('Table_S1_phase_breakpoints.csv').as_posix()} and the saved Figure 1 quantitative interpretation note.",
            "Table_S1_phase_breakpoints.csv",
        ),
        (
            "Table S2. Quantitative characteristics of phases",
            "This table reports the phase-level quantitative characteristics retained for Result 1, including total documents, average annual counts, growth rates, participating countries, and phase-specific peak years.",
            f"Derived from the existing Figure 1 quantitative workbook and retained in {TABLE_DIR.joinpath('Table_S2_quantitative_characteristics_of_phases.csv').as_posix()}.",
            "Table_S2_quantitative_characteristics_of_phases.csv",
        ),
        (
            "Table S3. SDG diffusion statistics",
            "This table combines yearly SDG diffusion breadth, diffusion intensity, and diffusion speed at the goal level. "
            "It preserves the saved yearly structure rather than collapsing to phase averages.",
            f"Derived from the latest non-archived SDG workbook and retained in {TABLE_DIR.joinpath('Table_S3_SDG_diffusion_statistics.csv').as_posix()}.",
            "Table_S3_SDG_diffusion_statistics.csv",
        ),
        (
            "Table S4. Concentration and inequality metrics",
            "This table consolidates yearly concentration metrics used for SDG inequality analysis, including HHI for count and impact where available.",
            f"Derived from the saved SDG concentration outputs and retained in {TABLE_DIR.joinpath('Table_S4_concentration_and_inequality_metrics.csv').as_posix()}.",
            "Table_S4_concentration_and_inequality_metrics.csv",
        ),
        (
            "Table S5. Ageing–energy SDG target weights (final v2)",
            "This table reports the retained final v2 target-level weighting framework, including the embedding base, the policy filter, the pre-normalization score, and the final normalized weight.",
            f"Copied from the validated final v2 target weight outputs and retained in {TABLE_DIR.joinpath('Table_S5_ageing_energy_target_weights.csv').as_posix()}.",
            "Table_S5_ageing_energy_target_weights.csv",
        ),
        (
            "Table S6. Pooled regression results",
            "This table reports the pooled regression rows saved for the ageing-energy alignment analysis.",
            f"Derived from the saved ageing_energy_story_analysis.xlsx outputs and retained in {TABLE_DIR.joinpath('Table_S6_pooled_regression_results.csv').as_posix()}.",
            "Table_S6_pooled_regression_results.csv",
        ),
        (
            "Table S7. Fixed-effects regression results",
            "This table isolates the currently saved fixed-effects row retained for the OADR-to-index model.",
            f"Derived from the saved ageing_energy_story_analysis.xlsx outputs and retained in {TABLE_DIR.joinpath('Table_S7_fixed_effects_regression_results.csv').as_posix()}.",
            "Table_S7_fixed_effects_regression_results.csv",
        ),
        (
            "Table S8. Policy influence and efficiency regressions",
            "This table reports the saved regressions linking the validated final v2 index to impact and efficiency outcomes.",
            f"Derived from the saved ageing_energy_story_analysis.xlsx outputs and retained in {TABLE_DIR.joinpath('Table_S8_policy_influence_and_efficiency_regressions.csv').as_posix()}.",
            "Table_S8_policy_influence_and_efficiency_regressions.csv",
        ),
        (
            "Table S9. Keyword taxonomy",
            "This table lists the retained clean keyword taxonomy for Result 3, including assigned category, mechanism relevance, and confidence.",
            f"Copied directly from the clean taxonomy mapping retained in {TABLE_DIR.joinpath('Table_S9_keyword_taxonomy.csv').as_posix()}.",
            "Table_S9_keyword_taxonomy.csv",
        ),
        (
            "Table S10. Topic share by period",
            "This table reports the window-level category counts and shares used for the Result 3 topic composition analyses.",
            f"Derived directly from the validated topic share output retained in {TABLE_DIR.joinpath('Table_S10_topic_share_by_period.csv').as_posix()}.",
            "Table_S10_topic_share_by_period.csv",
        ),
        (
            "Table S11. Policy attention elasticity results",
            "This table reports the retained Step 8 elasticity estimates, standard errors, p-values, and confidence intervals for the six taxonomy dimensions.",
            f"Copied directly from the saved Step 8 results retained in {TABLE_DIR.joinpath('Table_S11_policy_attention_elasticity.csv').as_posix()}.",
            "Table_S11_policy_attention_elasticity.csv",
        ),
        (
            "Table S12. Centrality comparison of key terms",
            "This table compares phase-wise semantic centrality, rank, and frequency for key macro and mechanism-related terms in the retained Step 9 network outputs.",
            f"Copied directly from the saved Step 9 comparison file retained in {TABLE_DIR.joinpath('Table_S12_centrality_comparison_of_key_terms.csv').as_posix()}.",
            "Table_S12_centrality_comparison_of_key_terms.csv",
        ),
        (
            "Table S13a. KL divergence by window",
            "This table reports the retained window-level KL divergence series under the equal-weight six-category reference distribution.",
            f"Copied directly from the saved Step 10 time-series file retained in {TABLE_DIR.joinpath('Table_S13a_KL_divergence_by_window.csv').as_posix()}.",
            "Table_S13a_KL_divergence_by_window.csv",
        ),
        (
            "Table S13b. KL divergence by phase",
            "This table reports the phase-averaged KL divergence values retained from Step 10.",
            f"Copied directly from the saved Step 10 phase summary file retained in {TABLE_DIR.joinpath('Table_S13b_KL_divergence_by_phase.csv').as_posix()}.",
            "Table_S13b_KL_divergence_by_phase.csv",
        ),
    ]
    for title, explanation, source_note, filename in table_meta:
        df = pd.read_csv(TABLE_DIR / filename)
        add_table_from_df(doc, title, explanation, source_note, df)

    doc.add_heading("Closing note", level=1)
    doc.add_paragraph(
        "This Word document is a consolidated SI draft assembled strictly from existing validated project outputs. "
        "No model specification, weighting rule, taxonomy definition, or regression design was changed during document assembly. "
        "Where source outputs were missing from the project tree, the missing item was not recreated and remains flagged in the supplementary manifest."
    )

    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    LEGACY_OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    doc.save(LEGACY_OUT_PATH)


if __name__ == "__main__":
    main()
