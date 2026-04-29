from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
STEP8_DIR = ROOT / "outputs" / "si" / "step8_pae"
STEP9_DIR = ROOT / "outputs" / "si" / "step9_network"
STEP10_DIR = ROOT / "outputs" / "si" / "step10_kl"
OUT_PATH = ROOT / "outputs" / "si" / "result3_si_analysis_figures.docx"


def set_default_font(document: Document, name: str = "Arial", size: int = 11) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style.font.size = Pt(size if style_name == "Normal" else size + (4 if style_name == "Title" else 2))


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True


def add_bullets_from_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)


def load_text(path: Path) -> list[str]:
    return path.read_text(encoding="utf-8").splitlines()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def main() -> None:
    step8_dir = STEP8_DIR
    step9_dir = STEP9_DIR
    step10_dir = STEP10_DIR
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    step8 = pd.read_csv(step8_dir / "step8_pae_ranked.csv")
    step9 = pd.read_csv(step9_dir / "step9_centrality_comparison.csv")
    step10 = pd.read_csv(step10_dir / "step10_kl_timeseries.csv")

    doc = Document()
    set_default_font(doc, "Arial", 11)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Supplementary Analyses for Part 3\nPersistent blind spots in ageing-related energy policy")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Step 8–10 results and figures compiled for SI use").italic = True

    doc.add_heading("Step 8. Policy Attention Elasticity", level=1)
    top = step8[["category", "epsilon_k", "p_value", "ci_low_95", "ci_high_95"]].copy()
    doc.add_paragraph(
        "This section reports log-log elasticity estimates linking thematic TopicShare to window-level old-age dependency ratios "
        "using the existing Step 7 topic-share results. No additional controls were used, and only windows with observed OADR values were retained."
    )
    add_bullets_from_lines(doc, [
        f"- Highest positive elasticity: {top.iloc[0]['category']} (epsilon = {top.iloc[0]['epsilon_k']:.3f}, 95% CI {top.iloc[0]['ci_low_95']:.3f} to {top.iloc[0]['ci_high_95']:.3f}, p = {top.iloc[0]['p_value']:.4f})",
        f"- Lowest elasticity: {top.iloc[-1]['category']} (epsilon = {top.iloc[-1]['epsilon_k']:.3f}, 95% CI {top.iloc[-1]['ci_low_95']:.3f} to {top.iloc[-1]['ci_high_95']:.3f}, p = {top.iloc[-1]['p_value']:.4f})",
    ])
    mech_rows = step8[step8["category"].isin([
        "housing and thermal comfort",
        "household behaviour and adoption",
        "affordability and income constraints",
        "energy transition and efficiency",
    ])]
    doc.add_paragraph("Mechanism-relevant themes of particular interest:")
    for row in mech_rows.itertuples(index=False):
        doc.add_paragraph(
            f"{row.category}: epsilon = {row.epsilon_k:.3f}, SE = {row.standard_error:.3f}, "
            f"95% CI [{row.ci_low_95:.3f}, {row.ci_high_95:.3f}], p = {row.p_value:.4f}.",
            style="List Bullet",
        )
    doc.add_picture(str(step8_dir / "figureS1_topic_elasticity_ranking.png"), width=Inches(6.8))
    add_caption(doc, "Figure S1. Topic elasticity ranking from Step 8. Points show elasticity estimates and horizontal lines show 95% confidence intervals.")
    add_bullets_from_lines(doc, load_text(step8_dir / "step8_pae_summary.md"))

    add_page_break(doc)

    doc.add_heading("Step 9. Network Path Dependency and Semantic Centrality", level=1)
    doc.add_paragraph(
        "This section summarises the dynamic keyword co-occurrence networks built for the four evolutionary phases. "
        "The comparison figure below is the clearest SI view of whether household-level mechanism terms remain peripheral relative to macro terms."
    )
    focus_terms = step9[step9["keyword"].isin([
        "health", "social", "care", "climate change", "accessibility",
        "affordable housing", "household", "thermal comfort",
        "energy poverty", "fuel poverty", "retrofit"
    ])].copy()
    for term in ["household", "thermal comfort", "energy poverty", "fuel poverty", "retrofit", "affordable housing", "climate change"]:
        sub = focus_terms[focus_terms["keyword"] == term]
        present = sub[sub["present_in_phase_network"] == True]
        if present.empty:
            doc.add_paragraph(f"{term}: not observed as a phase-level central node in the current network results.", style="List Bullet")
        else:
            ranks = ", ".join(
                f"{row.phase} rank {int(row.centrality_rank)} (freq {int(row.frequency)})"
                for row in present.itertuples(index=False)
            )
            doc.add_paragraph(f"{term}: {ranks}.", style="List Bullet")
    doc.add_picture(str(step9_dir / "figureS2_centrality_comparison.png"), width=Inches(6.5))
    add_caption(doc, "Figure S2. Relative semantic centrality comparison for key macro and mechanism terms across the four phases.")

    doc.add_paragraph("Phase-specific network views:")
    for phase_name, file_name in [
        ("Phase 1 network", "phase1_network.png"),
        ("Phase 2 network", "phase2_network.png"),
        ("Phase 3 network", "phase3_network.png"),
        ("Phase 4 network", "phase4_network.png"),
    ]:
        doc.add_paragraph(phase_name)
        doc.add_picture(str(step9_dir / file_name), width=Inches(5.8))
        add_caption(doc, f"{phase_name}. Network plot generated from phase-specific keyword co-occurrence among classified taxonomy terms.")
    add_bullets_from_lines(doc, load_text(step9_dir / "step9_network_summary.md"))

    add_page_break(doc)

    doc.add_heading("Step 10. Quantifying Mechanism Mismatch via Informational Divergence", level=1)
    doc.add_paragraph(
        "This section reports KL divergence between the observed topic distribution and the default equal-weight six-category baseline. "
        "The figure tracks whether policy expansion reduced or sustained mechanism mismatch over time."
    )
    min_row = step10.loc[step10["kl_divergence"].idxmin()]
    max_row = step10.loc[step10["kl_divergence"].idxmax()]
    last_row = step10.iloc[-1]
    doc.add_paragraph(
        f"Observed KL divergence ranges from {step10['kl_divergence'].min():.3f} in {min_row['time_window']} "
        f"to {step10['kl_divergence'].max():.3f} in {max_row['time_window']}. "
        f"The latest window ({last_row['time_window']}) records KL = {last_row['kl_divergence']:.3f}.",
        style="List Bullet",
    )
    doc.add_picture(str(step10_dir / "figureS3_kl_divergence_trend.png"), width=Inches(6.8))
    add_caption(doc, "Figure S3. KL divergence trend over time from Step 10, with the rapid expansion period visually highlighted.")
    add_bullets_from_lines(doc, load_text(step10_dir / "step10_kl_summary.md"))

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.add_run("End of SI compilation for Step 8–10").italic = True

    doc.save(OUT_PATH)


if __name__ == "__main__":
    main()
